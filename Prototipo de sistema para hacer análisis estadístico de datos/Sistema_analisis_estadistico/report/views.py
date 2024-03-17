from django.http import HttpResponse, JsonResponse
from django.shortcuts import redirect, render
import json

from data_analysis.utils.data_analysis_utils import guardar_resultado_en_sesion

import base64
from io import BytesIO
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 


@csrf_exempt
@require_http_methods(["POST"])
def exportar_reporte(request):
    try:
        data = json.loads(request.body.decode('utf-8'))

        doc = Document()
        doc.add_heading(data['titulo'], 0)
        doc.add_paragraph(data['descripcion'])

        # Manejo de las tablas
        for tabla_obj in data['tablas']:
            try:
                tabla = tabla_obj['filas']
                titulo_tabla = tabla_obj['titulo']
                if tabla:

                    # Añade el título de la tabla como un párrafo antes de la tabla
                    doc.add_paragraph(titulo_tabla, style='Heading4')

                    # Asegúrate de que la tabla tiene al menos una fila y una columna
                    if len(tabla) > 0 and len(tabla[0]) > 0:
                        doc_table = doc.add_table(rows=1, cols=len(tabla[0]))
                        
                        # Establecer estilo de la tabla
                        doc_table.style = 'Table Grid'

                        # Añadir los encabezados de la tabla
                        for i, encabezado in enumerate(tabla[0]):
                            doc_table.cell(0, i).text = encabezado
                            # Centro el texto del encabezado y hago que esté en negritas
                            paragraph = doc_table.cell(0, i).paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.runs
                            run[0].bold = True

                        # Añadir las filas de datos a la tabla
                        for datos_fila in tabla[1:]:
                            row_cells = doc_table.add_row().cells
                            for i, valor_celda in enumerate(datos_fila):
                                row_cells[i].text = str(valor_celda)  # Convertimos a string por seguridad
                                # Centro el texto de la celda
                                paragraph = row_cells[i].paragraphs[0]
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
            except IndexError as e:
                print(f'Error al agregar tabla: {str(e)}')     

        # Manejo de las imágenes
        for grafica_base64 in data.get('graficas', []):
            try:
                image_data = base64.b64decode(grafica_base64.split(',')[1])
                image_stream = BytesIO(image_data)
                doc.add_picture(image_stream, width=Inches(6))
            except Exception as e:
                # Manejo general de errores para imágenes
                print(f'Error al agregar imagen: {str(e)}')

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        response = HttpResponse(doc_io.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Reporte_Análisis.docx"'
        return response

    except json.JSONDecodeError as e:
        return JsonResponse({'error': f'Error al decodificar JSON: {str(e)}'}, status=400)
    except Exception as e:
        return JsonResponse({'error': f'Error al procesar la solicitud: {str(e)}'}, status=500)


def reporte_regresion_linear(request, file_name):
    # Asumiendo que los datos del reporte se almacenan en la sesión como JSON
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    # Eliminar los datos del reporte de la sesión después de usarlos
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_regresion_linear.html', {
        'file_name': file_name,
        'reporte': reporte_data
    })

def reporte_regresion_logistica(request, file_name):
    # Asumiendo que los datos del reporte se almacenan en la sesión como JSON
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    # Eliminar los datos del reporte de la sesión después de usarlos
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_regresion_logistica.html', {
        'file_name': file_name,
        'reporte': reporte_data
    })


def reporte_estadistica_descriptiva(request, file_name):
    # Asumiendo que los datos del reporte se almacenan en la sesión como JSON
    resultados_tendencia_central = request.session.get('resultados_tendencia_central', {})
    resultados_variabilidad = request.session.get('resultados_variabilidad', {})
    resultados_distribucion_frecuencias = request.session.get('resultados_distribucion_frecuencias', {})
    
    # Eliminar los datos del reporte de la sesión después de usarlos
    del request.session['resultados_tendencia_central']
    del request.session['resultados_variabilidad']
    del request.session['resultados_distribucion_frecuencias']
    
    # Verificar qué tipo de análisis se realizó y renderizar el template correspondiente
    if resultados_tendencia_central:
        return render(request, 'report/reporte_tendencia_central.html', {
            'file_name': file_name,
            'resultados_tendencia_central': resultados_tendencia_central
        })
    elif resultados_variabilidad:
        return render(request, 'report/reporte_variabilidad.html', {
            'file_name': file_name,
            'resultados_variabilidad': resultados_variabilidad
        })
    elif resultados_distribucion_frecuencias:
        return render(request, 'report/reporte_frecuencias.html', {
            'file_name': file_name,
            'resultados_distribucion_frecuencias': resultados_distribucion_frecuencias
        })
    else:
        resultado = {
            'Errores': ['No se encontraron resultados para mostrar. Revise que se hayan seleccionado columnas para el análisis y tipo de análisis.']
        }
        guardar_resultado_en_sesion(request, resultado)
        return redirect('file_handler:revisar_csv', file_name=file_name)


def reporte_correlacion_pearson(request, file_name):
    # Asumiendo que los datos del reporte se almacenan en la sesión como JSON
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    # Eliminar los datos del reporte de la sesión después de usarlos
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_correlacion_pearson.html', {
        'file_name': file_name,
        'reporte_data': reporte_data  # Asegúrate de que el contexto y el template coincidan en los nombres de las variables
    })

def reporte_correlacion_spearman(request, file_name):
    # Asumiendo que los datos del reporte se almacenan en la sesión como JSON
    reporte_data_json = request.session.get('reporte_data', '{}')
    reporte_data = json.loads(reporte_data_json)
    
    # Eliminar los datos del reporte de la sesión después de usarlos
    del request.session['reporte_data']
    
    return render(request, 'report/reporte_correlacion_spearman.html', {
        'file_name': file_name,
        'reporte_data': reporte_data  # Asegúrate de que el contexto y el template coincidan en los nombres de las variables
    })